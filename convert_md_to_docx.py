import markdown
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from latex2mathml.converter import convert
import re
from docx.oxml.shared import qn
from docx.oxml import OxmlElement

def create_math_element(latex_formula):
    try:
        # Convert LaTeX to MathML
        mathml = convert(latex_formula)
        
        # Create the OMath element
        omath = OxmlElement('m:oMath')
        omath.set(qn('xmlns:m'), 'http://schemas.openxmlformats.org/officeDocument/2006/math')
        
        # Parse the MathML and add it to the OMath element
        mathml_element = OxmlElement.parse_xml(mathml)
        omath.append(mathml_element)
        
        return omath
    except Exception as e:
        print(f"Error converting formula: {latex_formula}")
        print(f"Error: {str(e)}")
        return None

def convert_md_to_docx(md_file, docx_file):
    # Read the markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Create a new Word document
    doc = Document()
    
    # Split content by lines to process each line separately
    lines = md_content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Handle headers
        if line.startswith('#'):
            level = len(re.match(r'^#+', line).group())
            text = line.lstrip('#').strip()
            doc.add_heading(text, level=level)
            continue
            
        # Handle math formulas
        if '$' in line:
            # Split the line by math delimiters
            parts = re.split(r'(\$[^$]+\$)', line)
            p = doc.add_paragraph()
            
            for part in parts:
                if part.startswith('$') and part.endswith('$'):
                    # This is a math formula
                    formula = part[1:-1]  # Remove $ signs
                    
                    # Try to create math element
                    math_element = create_math_element(formula)
                    if math_element:
                        run = p.add_run()
                        run._element.append(math_element)
                    else:
                        # Fallback to plain text if conversion fails
                        run = p.add_run(formula)
                        run.font.name = 'Cambria Math'
                else:
                    # Regular text
                    if part.strip():
                        p.add_run(part)
            continue
            
        # Handle regular text
        if line:
            p = doc.add_paragraph()
            p.add_run(line)
    
    # Save the document
    doc.save(docx_file)

# Convert the file
convert_md_to_docx('笔记本/东南综评题目.md', '笔记本/东南综评题目.docx') 